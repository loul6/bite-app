from flask import Flask, request, send_file, render_template_string, flash, redirect, url_for
from werkzeug.utils import secure_filename
import io
from PIL import Image
# Conversión PDF <-> texto
try:
    from pdfminer.high_level import extract_text
except Exception:
    extract_text = None
# Para generar PDF desde texto
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except Exception:
    canvas = None
# Para leer docx
try:
    import docx
except Exception:
    docx = None

app = Flask(__name__)
app.secret_key = 'dev-key-bite-prototype'  # usar secret real en producción

# HTML embebido con Bootstrap y funcionalidad de texto a voz
INDEX_HTML = '''
<!doctype html>
<html lang="es">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Bite - Conversor de archivos</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
  </head>
  <body class="bg-light">
    <div class="container py-5">
      <div class="text-center mb-4">
        <h1 class="display-5 fw-bold">Bite</h1>
        <p class="lead">Convierte archivos de forma rápida y fácil</p>
      </div>
      {% with messages = get_flashed_messages() %}
        {% if messages %}
          <div class="alert alert-warning text-center">{{ messages[0] }}</div>
        {% endif %}
      {% endwith %}
      <div class="card shadow-sm p-4">
        <form method="post" action="/convert" enctype="multipart/form-data" class="row g-3 align-items-end">
          <div class="col-md-6">
            <label for="file" class="form-label">Selecciona tu archivo</label>
            <input class="form-control" type="file" id="file" name="file" required data-speech="Selecciona tu archivo a convertir">
          </div>
          <div class="col-md-4">
            <label for="target" class="form-label">Tipo de conversión</label>
            <select id="target" name="target" class="form-select" required data-speech="Selecciona el tipo de conversión">
              <option value="txt->pdf" data-speech="Convertir de texto a PDF">.txt → .pdf</option>
              <option value="pdf->txt" data-speech="Convertir de PDF a texto">.pdf → .txt</option>
              <option value="docx->txt" data-speech="Convertir de Word a texto">.docx → .txt</option>
              <option value="docx->pdf" data-speech="Convertir de Word a PDF">.docx → .pdf</option>
              <option value="img->png" data-speech="Convertir imagen a PNG">Imagen → .png</option>
              <option value="img->jpg" data-speech="Convertir imagen a JPG">Imagen → .jpg</option>
              <option value="img->webp" data-speech="Convertir imagen a WebP">Imagen → .webp</option>
            </select>
          </div>
          <div class="col-md-2 d-grid">
            <button type="submit" class="btn btn-primary btn-lg" data-speech="Convertir archivo">Convertir</button>
          </div>
        </form>
      </div>
      <footer class="mt-5 text-center text-muted small">
        Bite • Prototipo • Render deployment
      </footer>
    </div>
    <script>
      // Función para hablar
      function speak(text) {
        if ('speechSynthesis' in window) {
          const utterance = new SpeechSynthesisUtterance(text);
          utterance.lang = 'es-ES'; // español
          window.speechSynthesis.speak(utterance);
        }
      }
      // Escuchar cuando el mouse pasa sobre cualquier elemento con data-speech
      document.addEventListener('DOMContentLoaded', () => {
        const elements = document.querySelectorAll('[data-speech]');
        elements.forEach(el => {
          el.addEventListener('mouseover', () => {
            speak(el.getAttribute('data-speech'));
          });
        });
      });
    </script>
  </body>
</html>
'''

ALLOWED_IMAGE_EXT = {'png', 'jpg', 'jpeg', 'webp'}

@app.route('/')
def index():
    return render_template_string(INDEX_HTML)

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        flash('No se subió ningún archivo.')
        return redirect(url_for('index'))
    f = request.files['file']
    if f.filename == '':
        flash('Nombre de archivo vacío.')
        return redirect(url_for('index'))
    target = request.form.get('target')
    filename = secure_filename(f.filename)
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    data = f.read()
    file_stream = io.BytesIO(data)
    
    # txt -> pdf
    if target == 'txt->pdf':
        if ext != 'txt':
            flash('Para txt->pdf sube un .txt')
            return redirect(url_for('index'))
        if canvas is None:
            flash('Falta dependencia reportlab.')
            return redirect(url_for('index'))
        out = txt_to_pdf(file_stream)
        out.seek(0)
        return send_file(out, as_attachment=True, download_name=filename.rsplit('.', 1)[0] + '.pdf', mimetype='application/pdf')
    
    # pdf -> txt
    if target == 'pdf->txt':
        if ext != 'pdf':
            flash('Para pdf->txt sube un .pdf')
            return redirect(url_for('index'))
        if extract_text is None:
            flash('Falta dependencia pdfminer.six.')
            return redirect(url_for('index'))
        out = pdf_to_txt(file_stream)
        return send_file(io.BytesIO(out.encode('utf-8')), as_attachment=True, download_name=filename.rsplit('.', 1)[0] + '.txt', mimetype='text/plain')
    
    # docx -> txt
    if target == 'docx->txt':
        if ext != 'docx':
            flash('Para docx->txt sube un .docx')
            return redirect(url_for('index'))
        if docx is None:
            flash('Falta dependencia python-docx.')
            return redirect(url_for('index'))
        out = docx_to_txt(file_stream)
        return send_file(io.BytesIO(out.encode('utf-8')), as_attachment=True, download_name=filename.rsplit('.', 1)[0] + '.txt', mimetype='text/plain')
    
    # docx -> pdf
    if target == 'docx->pdf':
        if ext != 'docx':
            flash('Para docx->pdf sube un .docx')
            return redirect(url_for('index'))
        if docx is None or canvas is None:
            flash('Faltan dependencias: reportlab y python-docx')
            return redirect(url_for('index'))
        out = docx_to_pdf(file_stream)
        out.seek(0)
        return send_file(out, as_attachment=True, download_name=filename.rsplit('.', 1)[0] + '.pdf', mimetype='application/pdf')
    
    # image conversions
    if target.startswith('img->'):
        if ext not in ALLOWED_IMAGE_EXT:
            flash('Sube una imagen (png, jpg, jpeg, webp).')
            return redirect(url_for('index'))
        target_ext = target.split('->')[1]
        try:
            out = image_convert(file_stream, target_ext)
            out.seek(0)
            mimetype = 'image/' + ('jpeg' if target_ext == 'jpg' else target_ext)
            return send_file(out, as_attachment=True, download_name=filename.rsplit('.', 1)[0] + '.' + target_ext, mimetype=mimetype)
        except Exception as e:
            flash('Error en conversión de imagen: ' + str(e))
            return redirect(url_for('index'))
    
    flash('Conversión no soportada o datos inválidos.')
    return redirect(url_for('index'))

# ---- funciones de conversión ----
def txt_to_pdf(file_stream):
    file_stream.seek(0)
    text = file_stream.read().decode('utf-8', errors='replace')
    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=letter)
    width, height = letter
    margin_x = 40
    margin_y = 40
    lines = []
    for paragraph in text.splitlines():
        while len(paragraph) > 120:
            lines.append(paragraph[:120])
            paragraph = paragraph[120:]
        lines.append(paragraph)
    y = height - margin_y
    for line in lines:
        if y < margin_y:
            c.showPage()
            y = height - margin_y
        c.drawString(margin_x, y, line)
        y -= 12
    c.save()
    out.seek(0)
    return out

def pdf_to_txt(file_stream):
    file_stream.seek(0)
    return extract_text(file_stream)

def docx_to_txt(file_stream):
    file_stream.seek(0)
    doc = docx.Document(file_stream)
    paragraphs = [p.text for p in doc.paragraphs]
    return '\n'.join(paragraphs)

def image_convert(file_stream, target_ext):
    file_stream.seek(0)
    im = Image.open(file_stream)
    if target_ext in ('jpg', 'jpeg') and im.mode in ('RGBA', 'LA'):
        background = Image.new('RGB', im.size, (255, 255, 255))
        background.paste(im, mask=im.split()[-1])
        im = background
    elif target_ext in ('jpg', 'jpeg'):
        im = im.convert('RGB')
    out = io.BytesIO()
    pil_format = 'JPEG' if target_ext in ('jpg', 'jpeg') else target_ext.upper()
    im.save(out, format=pil_format)
    out.seek(0)
    return out

def docx_to_pdf(file_stream):
    """Convierte un .docx a PDF simple."""
    file_stream.seek(0)
    doc = docx.Document(file_stream)
    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=letter)
    width, height = letter
    margin_x = 40
    margin_y = 40
    y = height - margin_y
    for p in doc.paragraphs:
        text = p.text
        while len(text) > 120:
            c.drawString(margin_x, y, text[:120])
            y -= 12
            text = text[120:]
        if y < margin_y:
            c.showPage()
            y = height - margin_y
        c.drawString(margin_x, y, text)
        y -= 12
        if y < margin_y:
            c.showPage()
            y = height - margin_y
    c.save()
    out.seek(0)
    return out

if __name__ == '__main__':
    app.run(debug=True)
